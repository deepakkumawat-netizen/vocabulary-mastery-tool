import os
import sys
import json
import io
from pathlib import Path
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from groq import Groq
from dotenv import load_dotenv

sys.path.insert(0, str(Path(__file__).parent))
from database import (init_db, create_session, save_worksheet,
                      get_session_history, get_all_worksheets, save_rag_document)
from rag import rag_retriever
from nlp_adapter import get_grade_prompt_context, analyze_text_grade, GRADE_PROFILES, get_word_count
from mcp_tools import MCP_TOOLS, execute_mcp_tool
from security import (assert_public_url, read_upload_capped, client_ip,
                      generate_limiter, extract_limiter, upload_limiter)

load_dotenv()

GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_FALLBACK_MODELS = [
    GROQ_MODEL,
    "llama-3.1-8b-instant",
    "deepseek-r1-distill-llama-70b",
    "meta-llama/llama-4-scout-17b-16e-instruct",
]
_groq_client = None


def get_groq_client() -> Groq:
    global _groq_client
    if _groq_client is None:
        _groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    return _groq_client


def is_rate_limit_error(exc: Exception) -> bool:
    msg = str(exc).lower()
    return "429" in msg or "rate_limit" in msg or "rate limit" in msg or "tokens per day" in msg


@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        init_db()
    except Exception as e:
        print(f"[startup] DB init error: {e}")
    try:
        rag_retriever.build_index()
    except Exception as e:
        print(f"[startup] RAG build error: {e}")
    yield


app = FastAPI(title="Vocabulary Mastery Worksheet Tool", version="1.0.0", lifespan=lifespan)

# Allow this tool's own Render origin (same-origin SPA) plus localhost for
# dev. Wildcard origin used to be allowed, which meant anyone could
# cross-origin POST to the generate endpoint and drain the Groq quota.
# Override at deploy time with ALLOWED_ORIGINS="https://x.com,https://y.com"
# if you ever need to embed the API on another domain.
_default_origins = "https://vocabulary-mastery-tool.onrender.com,http://localhost:5173,http://localhost:5174,http://127.0.0.1:5173"
ALLOWED_ORIGINS = [o.strip() for o in os.getenv("ALLOWED_ORIGINS", _default_origins).split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)


# ── Cache-Control middleware ─────────────────────────────────────────────────
# The SPA's index.html references hashed asset filenames (index-{hash}.js).
# When we deploy a new build, the HTML changes to reference a NEW hash.
# Browsers that have the OLD HTML cached will keep loading the OLD JS bundle
# until their cache expires, even though Render is serving the new build —
# which is exactly the "user reports the same bug after we fix it" loop.
#
# Fix: serve HTML with no-cache (browser always re-checks) and hashed assets
# with long-immutable (safe — every build gets new hashes).
@app.middleware("http")
async def cache_control_headers(request, call_next):
    response = await call_next(request)
    path = request.url.path
    is_html = path == "/" or path.endswith(".html")
    is_hashed_asset = (
        path.startswith("/assets/")
        and any(path.endswith(ext) for ext in (".js", ".css", ".woff", ".woff2", ".ttf"))
    )
    if is_html:
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    elif is_hashed_asset:
        response.headers["Cache-Control"] = "public, max-age=31536000, immutable"
    return response


class WorksheetRequest(BaseModel):
    topic: str
    grade_level: int
    learning_objective: str
    board: Optional[str] = None  # CBSE, ICSE, RBSE, etc. — woven into the prompt for curriculum alignment
    source_text: Optional[str] = None
    additional_context: Optional[str] = None
    session_id: Optional[str] = None


class SessionCreate(BaseModel):
    metadata: Optional[dict] = None


class MCPToolCall(BaseModel):
    tool_name: str
    arguments: dict


class RAGDocRequest(BaseModel):
    content: str
    topic: Optional[str] = ""
    grade_level: Optional[int] = 0


# ── Sessions ──────────────────────────────────────────────────────────────────

@app.post("/api/sessions")
async def new_session(req: SessionCreate):
    session_id = create_session(req.metadata)
    return {"session_id": session_id}


@app.get("/api/sessions/{session_id}/history")
async def session_history(session_id: str):
    return {"session_id": session_id, "history": get_session_history(session_id)}


@app.get("/api/worksheets")
async def list_worksheets(limit: int = 20):
    return {"worksheets": get_all_worksheets(limit)}


# ── Generate ──────────────────────────────────────────────────────────────────

def _count_syllables(word: str) -> int:
    """Rough syllable counter — vowel groups, minus silent terminal 'e'."""
    import re as _re
    word = (word or "").lower().strip(".,!?;:'\"")
    if not word:
        return 1
    count = len(_re.findall(r'[aeiouy]+', word))
    if word.endswith('e') and count > 1:
        count -= 1
    return max(count, 1)


# Per-grade complexity caps for vocabulary words. The GRADE_PROFILES are
# instructive in the prompt but the model doesn't always follow them — at
# Grade 1 it will happily emit "photosynthesis" if the source PDF mentions
# it. This dict gives us a hard post-generation check so we can reject and
# retry with stricter instructions. Only applied to Grade 1-3 — older
# grades can handle whatever the LLM produces.
GRADE_VOCAB_CAPS = {
    1: {"max_syllables": 2, "max_chars": 7},
    2: {"max_syllables": 3, "max_chars": 8},
    3: {"max_syllables": 3, "max_chars": 9},
}


def _check_grade_complexity(data: dict, grade_level: int) -> "str | None":
    """For Grade 1-3, reject vocabulary words that are too long/complex.

    Returns an error string describing the offending words so the retry
    prompt can tell the model EXACTLY which words to swap out. Returning
    None means the generated content is grade-appropriate (or grade is
    high enough that we don't enforce this check).
    """
    caps = GRADE_VOCAB_CAPS.get(grade_level)
    if not caps:
        return None
    too_complex = []
    for vw in data.get("vocab_words", []):
        word = (vw.get("word") or "").strip().lower()
        if not word:
            continue
        if _count_syllables(word) > caps["max_syllables"] or len(word) > caps["max_chars"]:
            too_complex.append(word)
    if too_complex:
        return (
            f"Grade {grade_level} vocabulary must be at most {caps['max_syllables']} syllables and "
            f"{caps['max_chars']} letters per word. These words are too complex: "
            f"{', '.join(too_complex[:5])}. REPLACE each one with a simpler Grade {grade_level} word "
            f"on the same topic (1-2 syllable sight word, CVC pattern, decodable phonics)."
        )
    return None


def _validate_vocab(data: dict, min_items: int = 8) -> "str | None":
    """Return an error description string if the worksheet JSON is invalid, else None.
    `min_items` is the grade-calibrated minimum number of words/sentences/prompts."""
    vocab_words = data.get("vocab_words")
    if not isinstance(vocab_words, list) or len(vocab_words) < min_items:
        return f"vocab_words must have at least {min_items} items, got {len(vocab_words) if isinstance(vocab_words, list) else 'missing'}"

    matching = data.get("matching_section")
    if not isinstance(matching, dict) or not matching.get("items"):
        return "matching_section is missing or has no items"

    fib = data.get("fill_in_blank")
    if not isinstance(fib, dict):
        return "fill_in_blank section is missing"
    sentences = fib.get("sentences")
    if not isinstance(sentences, list) or len(sentences) < min_items:
        return f"fill_in_blank.sentences must have at least {min_items} items, got {len(sentences) if isinstance(sentences, list) else 'missing'}"

    sw = data.get("sentence_writing")
    if not isinstance(sw, dict) or not sw.get("prompts"):
        return "sentence_writing section is missing or has no prompts"
    prompts = sw.get("prompts")
    if not isinstance(prompts, list) or len(prompts) < min_items:
        return f"sentence_writing.prompts must have at least {min_items} items, got {len(prompts) if isinstance(prompts, list) else 'missing'}"

    return None


@app.post("/api/vocabulary/generate")
async def generate_worksheet(req: WorksheetRequest, request: Request):
    # Rate-limit the most-expensive endpoint per-IP — 15 generates / minute.
    # Without this, an attacker who finds the Render URL can drain the Groq
    # daily token quota in a loop with no auth in front of the API.
    await generate_limiter.check(client_ip(request))
    session_id = req.session_id or create_session()

    rag_retriever.build_index()
    rag_context = rag_retriever.build_context(
        f"{req.topic} grade {req.grade_level} {req.learning_objective}",
        grade_level=req.grade_level,
    )
    grade_ctx = get_grade_prompt_context(req.grade_level)

    # Pre-compute optional blocks to avoid backslashes inside f-string expressions
    source_block = (
        "\nSOURCE MATERIAL (MANDATORY — pick vocabulary words from THIS content; "
        "definitions, examples and context sentences must reflect what the source actually says):\n"
        f"---\n{req.source_text[:6000]}\n---\n"
    ) if req.source_text else ""
    additional_block = f"Additional Context: {req.additional_context}" if req.additional_context else ""
    rag_block = f"\n{rag_context}" if rag_context else ""
    # Curriculum board (CBSE / ICSE / RBSE / ...) becomes a hard constraint
    # in the prompt — vocabulary, examples, and exam style must match this
    # board's published syllabus for the selected grade.
    board_block = (
        f"\nCURRICULUM BOARD: {req.board}\n"
        f"All vocabulary, sentence examples, and instructional language must align with the "
        f"{req.board} syllabus and exam style for Grade {req.grade_level}. Use vocabulary and "
        f"context that an actual {req.board} Grade {req.grade_level} student would encounter in "
        f"their official textbooks and exams.\n"
    ) if req.board else ""
    ctx_block = f"{board_block}{source_block}{additional_block}\n{rag_block}".strip()

    def _build_prompt(extra_instructions: str = "") -> str:
        p = GRADE_PROFILES.get(req.grade_level, GRADE_PROFILES[7])
        n = get_word_count(req.grade_level)

        return f"""You are an expert educator and curriculum specialist.
Your task is to create a grade-calibrated Vocabulary Mastery Worksheet.

{grade_ctx}

CONTENT DETAILS:
Topic: {req.topic}
Learning Objective: {req.learning_objective}
{ctx_block}

CRITICAL RULES:
1. Generate EXACTLY {n} vocabulary words — this count is calibrated for Grade {req.grade_level} attention span. Do NOT generate more or fewer.
2. ALL {n} vocabulary words must be exactly right for Grade {req.grade_level} students (age {req.grade_level + 5}-{req.grade_level + 6}).
3. Every definition must use SIMPLER words than the target word — a Grade {req.grade_level} student must understand it.
4. Fill-in-blank sentences must be at Grade {req.grade_level} reading level: {p['sentence']}
5. Sentence writing hints must be Grade {req.grade_level} appropriate: {p['hint_style']}
6. Do NOT use words from other grade levels. Do NOT use placeholder words like 'word1'.
{"7. SOURCE MATERIAL is provided above and is the AUTHORITATIVE basis for this worksheet. EVERY vocabulary word MUST be picked from words that actually appear in the SOURCE MATERIAL — do NOT introduce vocabulary that is not present in the source. Definitions must reflect how the word is used in the source. Fill-in-blank sentences must mirror the source's topic/context." if req.source_text else ""}
{extra_instructions}

Return ONLY valid JSON. No markdown fences. No prose outside the JSON.

{{
  "vocab_words": [
    {{"word": "actual Grade {req.grade_level} word from topic", "definition": "Grade {req.grade_level}-appropriate definition", "part_of_speech": "noun|verb|adjective|adverb"}},
    ... EXACTLY {n} words total, all relevant to '{req.topic}', all appropriate for Grade {req.grade_level}
  ],
  "matching_section": {{
    "title": "Section 1: Match the Word to Its Meaning",
    "instructions": "Grade {req.grade_level}-appropriate matching instruction (1 sentence).",
    "items": [
      {{"word": "word from vocab_words", "definition": "Grade {req.grade_level}-appropriate definition"}},
      ... all {n} words, definitions in SHUFFLED order (not matching vocab_words order)
    ]
  }},
  "fill_in_blank": {{
    "title": "Section 2: Fill in the Blank",
    "instructions": "Grade {req.grade_level}-appropriate instruction (1 sentence).",
    "word_bank": ["all {n} vocab words listed here"],
    "sentences": [
      {{"sentence": "Grade {req.grade_level} sentence with ___ blank for the answer word.", "answer": "the correct vocab word"}},
      ... EXACTLY {n} sentences total, one for EACH vocabulary word
    ]
  }},
  "sentence_writing": {{
    "title": "Section 3: Write Your Own Sentences",
    "instructions": "Grade {req.grade_level}-appropriate writing instruction.",
    "prompts": [
      {{"word": "vocab word", "hint": "{p['hint_style']}", "example": "Grade {req.grade_level}-appropriate example sentence"}},
      ... EXACTLY {n} prompts total, one for EACH vocabulary word
    ]
  }}
}}"""

    def _sse(obj: dict) -> str:
        return f"data: {json.dumps(obj)}\n\n"

    def stream_gen():
        max_attempts = 5
        extra_instructions = ""
        last_reason = ""
        model_idx = 0

        for attempt in range(1, max_attempts + 1):
            if attempt > 1:
                yield _sse({"type": "retry", "attempt": attempt, "reason": last_reason})

            current_model = GROQ_FALLBACK_MODELS[min(model_idx, len(GROQ_FALLBACK_MODELS) - 1)]
            yield _sse({"type": "progress", "message": f"Attempt {attempt}: calling {current_model}…"})

            prompt = _build_prompt(extra_instructions)
            collected_chunks = []

            try:
                stream = get_groq_client().chat.completions.create(
                    model=current_model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=4500,
                    stream=True,
                )

                for chunk in stream:
                    delta = chunk.choices[0].delta.content or ""
                    if delta:
                        collected_chunks.append(delta)
                        yield _sse({"type": "token", "content": delta})

            except Exception as exc:
                last_reason = str(exc)
                if model_idx < len(GROQ_FALLBACK_MODELS) - 1:
                    model_idx += 1
                    next_model = GROQ_FALLBACK_MODELS[model_idx]
                    yield _sse({"type": "status", "message": f"Model error — switching to {next_model}…"})
                    extra_instructions = ""
                else:
                    extra_instructions = f"IMPORTANT: Fix the following error from the previous attempt: {last_reason}\n"
                continue

            raw = "".join(collected_chunks).strip()
            for fence in ("```json", "```"):
                if raw.startswith(fence):
                    raw = raw[len(fence):]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
            import re as _re
            raw = _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', raw)

            yield _sse({"type": "status", "message": "Parsing JSON response…"})

            # Isolate the JSON object (drop any stray prose before/after)
            first, last = raw.find("{"), raw.rfind("}")
            if first != -1 and last != -1 and last > first:
                raw = raw[first:last + 1]

            data = None
            try:
                data = json.loads(raw)
            except json.JSONDecodeError as exc:
                # LLMs often emit unescaped quotes/apostrophes or trailing commas.
                # json-repair fixes these instead of forcing another full retry.
                try:
                    from json_repair import repair_json
                    repaired = repair_json(raw)
                    data = json.loads(repaired)
                except Exception:
                    last_reason = f"Invalid JSON: {exc}"
                    extra_instructions = (
                        "CRITICAL: Your previous response was not valid JSON. "
                        "Return ONLY a raw JSON object — no markdown fences, no prose. "
                        "Escape every double-quote inside string values as \\\". "
                        "Do not put a comment line starting with '...' inside the JSON.\n"
                    )
                    continue

            yield _sse({"type": "status", "message": "Validating worksheet structure…"})

            validation_error = _validate_vocab(data, min_items=get_word_count(req.grade_level))
            if validation_error:
                last_reason = f"Validation failed: {validation_error}"
                _n = get_word_count(req.grade_level)
                extra_instructions = (
                    f"IMPORTANT: Fix this validation error from your previous attempt: {validation_error}. "
                    f"Ensure vocab_words has exactly {_n} items, matching_section has {_n} items, "
                    f"fill_in_blank has exactly {_n} sentences, and sentence_writing has exactly {_n} prompts.\n"
                )
                continue

            # For young grades, reject worksheets whose vocab words exceed
            # the syllable/length caps. The model often pulls complex words
            # straight from the source PDF (e.g. "photosynthesis" for a
            # Grade 1 student) even with the NLP calibration in the
            # prompt — this is the post-generation enforcement step.
            complexity_error = _check_grade_complexity(data, req.grade_level)
            if complexity_error:
                last_reason = f"Grade-complexity failed: {complexity_error}"
                yield _sse({"type": "status", "message": "Words too complex for the grade — regenerating with simpler vocabulary…"})
                extra_instructions = complexity_error + "\n"
                continue

            # All checks passed — save and emit complete event
            yield _sse({"type": "status", "message": "Saving worksheet…"})

            full_content = {
                **data,
                "rag_context_used": bool(rag_context),
            }

            try:
                worksheet_id = save_worksheet(
                    session_id=session_id,
                    topic=req.topic,
                    grade_level=req.grade_level,
                    learning_objective=req.learning_objective,
                    content=full_content,
                )

                save_rag_document(
                    content=(
                        f"vocabulary worksheet topic {req.topic} grade {req.grade_level} "
                        f"objective {req.learning_objective} words "
                        + " ".join(w["word"] for w in data.get("vocab_words", []))
                    ),
                    doc_type="worksheet",
                    topic=req.topic,
                    grade_level=req.grade_level,
                )
                rag_retriever.build_index()
            except Exception as exc:
                yield _sse({"type": "error", "message": f"Database error: {exc}"})
                return

            yield _sse({
                "type": "complete",
                "session_id": session_id,
                "worksheet_id": worksheet_id,
                "worksheet": full_content,
            })
            return

        # Exhausted all retries
        yield _sse({"type": "error", "message": f"Failed after {max_attempts} attempts. Last error: {last_reason}"})

    return StreamingResponse(
        stream_gen(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# ── Export DOCX ───────────────────────────────────────────────────────────────

@app.post("/api/vocabulary/export/docx")
async def export_docx(payload: dict):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ws = payload.get("worksheet", {})
    topic = payload.get("topic", "Vocabulary")
    grade = payload.get("grade_level", "")
    objective = payload.get("learning_objective", "")

    doc = Document()

    title = doc.add_heading("Vocabulary Mastery Worksheet", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Topic: {topic}  |  Grade: {grade}  |  Objective: {objective}")
    doc.add_paragraph("Name: ____________________________   Date: _______________")
    doc.add_paragraph()

    matching = ws.get("matching_section", {})
    if matching:
        doc.add_heading(matching.get("title", "Section 1: Matching"), 1)
        doc.add_paragraph(matching.get("instructions", ""))
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Vocabulary Word"
        hdr[1].text = "Definition"
        for item in matching.get("items", []):
            row = tbl.add_row().cells
            row[0].text = item.get("word", "")
            row[1].text = ""
        doc.add_paragraph()

    fib = ws.get("fill_in_blank", {})
    if fib:
        doc.add_heading(fib.get("title", "Section 2: Fill in the Blank"), 1)
        doc.add_paragraph(fib.get("instructions", ""))
        wb = ", ".join(fib.get("word_bank", []))
        doc.add_paragraph(f"Word Bank: [ {wb} ]")
        doc.add_paragraph()
        for i, s in enumerate(fib.get("sentences", []), 1):
            doc.add_paragraph(f"{i}. {s.get('sentence', '')}")
        doc.add_paragraph()

    sw = ws.get("sentence_writing", {})
    if sw:
        doc.add_heading(sw.get("title", "Section 3: Write Your Own Sentences"), 1)
        doc.add_paragraph(sw.get("instructions", ""))
        doc.add_paragraph()
        for i, p in enumerate(sw.get("prompts", []), 1):
            doc.add_paragraph(f"{i}. Word: {p.get('word', '')}")
            doc.add_paragraph(f"   Hint: {p.get('hint', '')}")
            doc.add_paragraph("   My sentence: _________________________________________________")
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="vocabulary_{topic}.docx"'},
    )


# ── RAG document upload ───────────────────────────────────────────────────────

@app.post("/api/rag/add-text")
async def add_rag_text(req: RAGDocRequest):
    doc_id = save_rag_document(req.content, "knowledge", req.topic, req.grade_level)
    rag_retriever.build_index()
    return {"success": True, "doc_id": doc_id}


@app.post("/api/rag/add-file")
async def add_rag_file(request: Request, file: UploadFile = File(...)):
    await upload_limiter.check(client_ip(request))
    # Stream-read with a 10MB cap. Previously `await file.read()` buffered the
    # entire upload in memory before any size check — a 2GB POST would OOM
    # the Render container before the endpoint could even reject it.
    raw = await read_upload_capped(file)
    content = ""
    if file.filename.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(raw))
        content = " ".join(p.extract_text() or "" for p in reader.pages)
    elif file.filename.endswith(".docx"):
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(raw))
        content = " ".join(p.text for p in doc.paragraphs)
    else:
        content = raw.decode("utf-8", errors="ignore")

    content = (content or "").strip()
    doc_id = save_rag_document(content[:6000], "file", file.filename, 0)
    rag_retriever.build_index()
    # Return the extracted text so the frontend can pass it as source_text.
    return {
        "success": True,
        "doc_id": doc_id,
        "chars_indexed": len(content),
        "text": content[:8000],
        "filename": file.filename,
    }


# ── Hero image proxy (NanoBanana AI) ─────────────────────────────────────────
# Frontend's Landing page asks for `/api/hero-image?seed=N` and gets back a
# fresh cartoon JPEG generated by NanoBanana. The secret key lives in
# NANOBANANA_API_KEY on the server, never reaches the browser.
#
# NanoBanana is async — POST a generate task, poll record-info until done,
# then fetch the result image. Typical end-to-end takes ~15s.

VOCAB_HERO_PROMPTS = [
    "3D Pixar cartoon of a smiling kid holding an open dictionary with glowing alphabet letters floating around them, bright cheerful colors, clean white background, educational illustration",
    "3D Pixar cartoon of colorful ABC alphabet blocks stacked on a wooden desk with floating vocabulary words and a friendly pencil character, clean white background, educational",
    "3D Pixar cartoon of a happy student with backpack holding stack of vocabulary flashcards with words like JOY HOPE WISE floating in air, bright colors, clean white background",
    "3D Pixar cartoon of a giant magnifying glass examining a dictionary page with words leaping off the page in colorful 3D letters, clean white background, educational",
    "3D Pixar cartoon of a friendly book character with arms and a smile teaching glowing speech bubbles with vocabulary words, classroom setting, clean white background",
    "3D Pixar cartoon of a cute owl wearing graduation cap holding a dictionary surrounded by floating alphabet letters and stars, bright cheerful colors, clean white background",
]


@app.get("/api/hero-image")
async def hero_image(request: Request, seed: Optional[int] = None):
    """Proxy to Hugging Face FLUX.1-schnell — fresh cartoon hero image per
    visit. The HF token lives in HF_API_TOKEN on the server.

    Total time: ~5-10s warm, up to ~60s on cold start. The `x-wait-for-model`
    header makes HF block until the model is loaded instead of returning
    503 — better UX than the alternative (frontend shows gradient longer
    but always gets an image)."""
    await extract_limiter.check(client_ip(request))
    hf_token = (os.getenv("HF_API_TOKEN") or "").strip()
    if not hf_token:
        raise HTTPException(status_code=503, detail="HF_API_TOKEN not configured")

    import random as _random
    if seed is None:
        seed = _random.randint(1, 999999)
    prompt = VOCAB_HERO_PROMPTS[seed % len(VOCAB_HERO_PROMPTS)]

    try:
        import httpx
        async with httpx.AsyncClient(timeout=120.0) as cx:
            r = await cx.post(
                "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-schnell",
                headers={
                    "Authorization": f"Bearer {hf_token}",
                    "Content-Type": "application/json",
                    "x-wait-for-model": "true",
                    "Accept": "image/png",
                },
                json={"inputs": prompt, "parameters": {"seed": seed, "num_inference_steps": 4}},
            )
        if r.status_code == 503:
            # Model still loading — surface clearly so frontend can show
            # the gradient placeholder and the user can refresh in a bit.
            raise HTTPException(status_code=503, detail="HF model is warming up; try again in 30s")
        if r.status_code == 429:
            raise HTTPException(status_code=429, detail="HF free-tier rate limit hit")
        if r.status_code != 200:
            # JSON error responses come through here
            detail = r.text[:200] if r.text else f"HTTP {r.status_code}"
            raise HTTPException(status_code=502, detail=f"HF returned {r.status_code}: {detail}")
        from fastapi.responses import Response
        return Response(
            content=r.content,
            media_type=r.headers.get("content-type", "image/png"),
            # 1-hour cache — keeps repeated visits cheap.
            headers={"Cache-Control": "public, max-age=3600"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Hero image generation failed: {e}")


@app.post("/api/extract-url")
async def extract_url(req: dict, request: Request):
    """Fetch a webpage and return cleaned text as source_text."""
    await extract_limiter.check(client_ip(request))
    # SSRF guard: reject loopback, RFC1918, AWS metadata, IPv6 ULA, etc.
    # Without this, a teacher (or a curl user) could submit
    # http://169.254.169.254/latest/meta-data/iam and exfiltrate the
    # Render/AWS service identity from the container.
    url = assert_public_url((req.get("url") or "").strip())
    try:
        import httpx, re as _re
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=20.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; VocabularyTool/1.0)"}) as cx:
            r = await cx.get(url)
            # Re-validate the FINAL URL after redirects so an attacker can't
            # bounce through a public host into a private one.
            assert_public_url(str(r.url))
            r.raise_for_status()
            soup = BeautifulSoup(r.text, "html.parser")
            for bad in soup(["script", "style", "noscript", "iframe", "nav", "footer", "header", "form", "aside"]):
                bad.decompose()
            title = (soup.title.string or "").strip() if soup.title else ""
            main = soup.find("main") or soup.find("article") or soup.body or soup
            text = _re.sub(r"\s+\n", "\n", main.get_text("\n", strip=True))
            text = _re.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Could not extract readable text from this page.")
        return {"success": True, "title": title, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except httpx.HTTPStatusError as e:
        # Surface upstream 4xx as 4xx (not a generic 500) so the UI can
        # tell the teacher "that page returned 404", not "server error".
        raise HTTPException(status_code=e.response.status_code,
                            detail=f"URL fetch failed: {e.response.status_code} {e.response.reason_phrase}")
    except httpx.RequestError as e:
        raise HTTPException(status_code=502, detail=f"Could not reach URL: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"URL fetch failed: {e}")


@app.post("/api/auto-fields")
async def auto_fields(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    """Read uploaded text and propose a Topic + Learning Objective + grade.

    Saves the teacher from typing — they upload a PDF / URL / YouTube
    transcript and we suggest what to put in the Topic and Learning
    Objective fields, calibrated for the grade if they've picked one.
    """
    source_text = (req.get("source_text") or "").strip()
    if not source_text:
        raise HTTPException(status_code=400, detail="source_text is required.")
    grade = req.get("grade_level")
    board = (req.get("board") or "").strip()
    grade_hint = f"Calibrate the topic + objective for Grade {grade} students. " if grade else ""
    board_hint = f"This is for the {board} curriculum — phrase the topic and objective to match {board} terminology and exam style. " if board else ""

    prompt = (
        "You are an expert vocabulary curriculum designer.\n"
        "Read the SOURCE MATERIAL below and propose a clean, classroom-ready\n"
        "  • Topic (a short noun phrase, 4-10 words, naming the main idea/theme of the material)\n"
        "  • Learning Objective (one sentence starting with 'Students will…' "
        "describing what vocabulary skill the student will gain).\n"
        f"{grade_hint}{board_hint}"
        "Return ONLY a JSON object with keys 'topic' and 'learning_objective'. "
        "No markdown, no prose outside the JSON.\n\n"
        "SOURCE MATERIAL:\n---\n"
        f"{source_text[:6000]}\n---\n\n"
        'JSON: {"topic": "...", "learning_objective": "Students will ..."}'
    )

    try:
        client = get_groq_client()
        completion = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You return strict JSON. No markdown fences."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=400,
            response_format={"type": "json_object"},
        )
        import json as _json
        raw = completion.choices[0].message.content.strip()
        data = _json.loads(raw)
        topic = (data.get("topic") or "").strip()
        objective = (data.get("learning_objective") or "").strip()
        if not topic and not objective:
            raise HTTPException(status_code=502, detail="Model returned no fields.")
        return {"success": True, "topic": topic, "learning_objective": objective}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"auto-fields failed: {e}")


async def _youtube_metadata_fallback(video_id: str, url: str) -> dict | None:
    """Build source_text from public YouTube metadata when the transcript API
    is IP-blocked. Tries three sources in order so we always return
    *something* the LLM can use:

      1. YouTube oEmbed (always works, gives title + channel)
      2. noembed.com aggregator (often includes a longer description)
      3. youtube.com/watch with consent cookie (full shortDescription)
    """
    import httpx, re as _re2
    title = ""
    author = ""
    description = ""

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; VocabularyTool/1.0)"}) as cx:
            r = await cx.get("https://www.youtube.com/oembed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}", "format": "json"})
            if r.status_code == 200:
                j = r.json()
                title = (j.get("title") or "").strip()
                author = (j.get("author_name") or "").strip()
    except Exception as _e:
        print(f"[YouTube oEmbed] failed: {_e}")

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; VocabularyTool/1.0)"}) as cx:
            r = await cx.get("https://noembed.com/embed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}"})
            if r.status_code == 200:
                j = r.json()
                if not title:  title  = (j.get("title") or "").strip()
                if not author: author = (j.get("author_name") or "").strip()
                description = (j.get("description") or "").strip()
    except Exception as _e:
        print(f"[YouTube noembed] failed: {_e}")

    try:
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                              "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept-Language": "en-US,en;q=0.9",
                "Cookie": "CONSENT=YES+cb.20210328-17-p0.en+FX+000",
            }) as cx:
            r = await cx.get(f"https://www.youtube.com/watch?v={video_id}")
            if r.status_code == 200:
                soup = BeautifulSoup(r.text, "html.parser")
                if not title:
                    ot = soup.find("meta", attrs={"property": "og:title"})
                    if ot: title = (ot.get("content") or "").strip()
                ot_desc = soup.find("meta", attrs={"property": "og:description"})
                if ot_desc and not description:
                    description = (ot_desc.get("content") or "").strip()
                for s in soup.find_all("script"):
                    txt = s.string or ""
                    if "shortDescription" in txt:
                        mm = _re2.search(r'"shortDescription":"((?:\\.|[^"\\])*)"', txt)
                        if mm:
                            full = mm.group(1).encode("utf-8").decode("unicode_escape", errors="ignore")
                            if len(full) > len(description):
                                description = full
                            break
    except Exception as _e:
        print(f"[YouTube watch-page] failed: {_e}")

    pieces = []
    if title:       pieces.append(f"Video title: {title}")
    if author:      pieces.append(f"Channel: {author}")
    if description: pieces.append(f"\n{description}")
    text = "\n".join(pieces).strip()
    if not text or len(text) < 20:
        return None
    return {"title": title or f"YouTube video {video_id}", "text": text}


@app.post("/api/extract-youtube")
async def extract_youtube(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    """Fetch a YouTube transcript and return it as source_text. If YouTube
    blocks the transcript API (typical on cloud IPs), fall back to scraping
    the video's title + description."""
    url = (req.get("url") or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="url is required")
    import re as _re
    m = _re.search(r"(?:v=|youtu\.be/|/embed/|/shorts/)([A-Za-z0-9_-]{11})", url)
    if not m:
        raise HTTPException(status_code=400, detail="Could not detect a YouTube video id in that URL.")
    video_id = m.group(1)
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        proxy_config = None
        wh_user = os.getenv("WEBSHARE_PROXY_USERNAME")
        wh_pass = os.getenv("WEBSHARE_PROXY_PASSWORD")
        if wh_user and wh_pass:
            try:
                from youtube_transcript_api.proxies import WebshareProxyConfig
                proxy_config = WebshareProxyConfig(proxy_username=wh_user, proxy_password=wh_pass)
            except Exception as _e:
                print(f"[YouTube] Webshare proxy import failed: {_e}")

        if proxy_config is not None:
            transcript = YouTubeTranscriptApi(proxy_config=proxy_config).fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        elif hasattr(YouTubeTranscriptApi, "get_transcript"):
            chunks = YouTubeTranscriptApi.get_transcript(video_id)
            text = " ".join((c.get("text", "") if isinstance(c, dict) else getattr(c, "text", "")) for c in chunks).strip()
        else:
            transcript = YouTubeTranscriptApi().fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Transcript was empty.")
        return {"success": True, "video_id": video_id, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except Exception as e:
        msg = str(e)
        blocked = ("blocking requests" in msg.lower()
                   or ("ip" in msg.lower() and "block" in msg.lower())
                   or "could not retrieve a transcript" in msg.lower())

        if blocked:
            fallback = await _youtube_metadata_fallback(video_id, url)
            if fallback:
                return {
                    "success": True,
                    "video_id": video_id,
                    "url": url,
                    "text": fallback["text"][:8000],
                    "chars": len(fallback["text"]),
                    "title": fallback["title"],
                    "note": "Transcript was blocked from our server's IP — using the video's title + description instead. For best results, paste the full transcript manually.",
                }

        raise HTTPException(
            status_code=502 if blocked else 500,
            detail=(
                "YouTube blocks transcript requests from our cloud server's IP, and we "
                "couldn't reach the video's public description either. Please open the "
                "video in your browser → click the ⋯ menu (or 'Show transcript' in the "
                "description) → copy the transcript → paste it into the textarea below."
                if blocked else f"YouTube transcript fetch failed: {msg}"
            ),
        )


# ── MCP Tools ─────────────────────────────────────────────────────────────────

@app.get("/mcp/tools")
async def list_mcp_tools():
    return {"tools": MCP_TOOLS}


@app.post("/mcp/tools/call")
async def call_mcp_tool(req: MCPToolCall):
    try:
        result = await execute_mcp_tool(req.tool_name, req.arguments)
        return {"success": True, "result": result}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health():
    return {"status": "ok", "tool": "vocabulary-mastery-worksheet", "model": GROQ_MODEL}


# ── Serve frontend ────────────────────────────────────────────────────────────

frontend_dir = Path(__file__).parent.parent / "frontend" / "dist"

if frontend_dir.exists():
    from fastapi.responses import FileResponse

    @app.get("/")
    async def serve_index():
        return FileResponse(
            str(frontend_dir / "index.html"),
            headers={"Cache-Control": "no-store, no-cache, must-revalidate"},
        )

    @app.get("/assets/index.js")
    async def serve_js():
        return FileResponse(
            str(frontend_dir / "assets" / "index.js"),
            media_type="application/javascript",
            headers={"Cache-Control": "no-cache, must-revalidate"},
        )

    try:
        app.mount("/", StaticFiles(directory=str(frontend_dir), html=True), name="static")
    except Exception as e:
        print(f"[startup] Static files mount skipped: {e}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002, reload=True)
